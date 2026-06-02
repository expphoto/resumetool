"""DOCX rendering for structured resumes.

Builds an ATS-friendly Word document with the standard section order:
header, summary, skills, experience, education, certifications. Uses only
default Word built-in styles so the output renders consistently across
Word, LibreOffice, and Google Docs.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from resumetool.types import Resume


def render_docx(
    structured_resume_path: str | Path,
    output_path: str | Path | None = None,
) -> str:
    """Render a structured resume JSON to a DOCX file.

    Returns the absolute path of the rendered file.
    """
    src = Path(structured_resume_path)
    if not src.exists():
        raise FileNotFoundError(f"Resume JSON not found: {src}")
    resume = Resume.model_validate(json.loads(src.read_text(encoding="utf-8")))

    out = Path(output_path) if output_path else src.with_suffix(".docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = _new_document()
    _write_header(doc, resume)
    if resume.summary:
        _write_section(doc, "Summary")
        doc.add_paragraph(resume.summary)
    if resume.skills:
        _write_section(doc, "Skills")
        _write_skills(doc, resume.skills)
    if resume.experience:
        _write_section(doc, "Experience")
        for job in resume.experience:
            _write_job(doc, job)
    if resume.education:
        _write_section(doc, "Education")
        for line in resume.education:
            doc.add_paragraph(line, style="List Bullet")
    if resume.certifications:
        _write_section(doc, "Certifications")
        for line in resume.certifications:
            doc.add_paragraph(line, style="List Bullet")

    if resume.contact_info_redacted:
        note = doc.add_paragraph()
        run = note.add_run("Contact information redacted for privacy.")
        run.italic = True
        run.font.size = Pt(9)

    doc.save(str(out))
    return str(out.resolve())


def _new_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    sec = doc.sections[0]
    sec.top_margin = Pt(48)
    sec.bottom_margin = Pt(48)
    sec.left_margin = Pt(54)
    sec.right_margin = Pt(54)
    return doc


def _write_header(doc: Document, resume: Resume) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(resume.full_name or "Candidate")
    run.bold = True
    run.font.size = Pt(20)
    if not resume.contact_info_redacted and resume.full_name:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(resume.full_name)
        sub_run.font.size = Pt(10)


def _write_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)


def _write_skills(doc: Document, skills: list[str]) -> None:
    p = doc.add_paragraph()
    p.add_run(", ".join(skills))


def _write_job(doc: Document, job) -> None:
    head = doc.add_paragraph()
    title_run = head.add_run(f"{job.title} — {job.company}")
    title_run.bold = True
    head.add_run("\t")
    dur_run = head.add_run(job.duration)
    dur_run.italic = True
    for bullet in job.bullet_points:
        doc.add_paragraph(bullet, style="List Bullet")
