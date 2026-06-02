"""Unit tests for the resume parsing layer."""
from pathlib import Path

import pytest

from resumetool.parsers import parse_any, parse_docx, parse_pdf, parse_txt
from resumetool.parsers.pdf import parse_pdf as _parse_pdf


FIXTURES = Path(__file__).parent / "fixtures"


def test_parse_txt_round_trip(tmp_path):
    src = tmp_path / "r.txt"
    src.write_text("Hello world\nLine two")
    assert parse_txt(src) == "Hello world\nLine two"


def test_parse_any_dispatches_to_txt(tmp_path):
    src = tmp_path / "r.md"
    src.write_text("markdown content")
    assert parse_any(src) == "markdown content"


def test_parse_any_dispatches_to_docx(tmp_path):
    from docx import Document

    d = Document()
    d.add_paragraph("Header line")
    d.add_paragraph("Body line", style="List Bullet")
    p = tmp_path / "r.docx"
    d.save(p)
    out = parse_any(p)
    assert "Header line" in out
    assert "Body line" in out


def test_parse_any_rejects_unknown_extension(tmp_path):
    p = tmp_path / "r.xyz"
    p.write_text("x")
    with pytest.raises(ValueError):
        parse_any(p)


def test_parse_docx_collects_paragraphs_and_bullets(tmp_path):
    from docx import Document

    d = Document()
    d.add_paragraph("Title")
    d.add_paragraph("Bullet one", style="List Bullet")
    d.add_paragraph("Bullet two", style="List Bullet")
    p = tmp_path / "r.docx"
    d.save(p)
    out = parse_docx(p)
    assert "Title" in out
    assert "Bullet one" in out
    assert "Bullet two" in out
    # Bullets should be indented
    bullet_lines = [line for line in out.splitlines() if line.startswith("  - ")]
    assert len(bullet_lines) == 2


def test_parse_docx_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        parse_docx(tmp_path / "missing.docx")


def test_parse_docx_wrong_extension_raises(tmp_path):
    p = tmp_path / "r.txt"
    p.write_text("x")
    with pytest.raises(ValueError):
        parse_docx(p)


def test_parse_pdf_empty_pdf_returns_empty_string(tmp_path):
    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=612, height=792)
    p = tmp_path / "empty.pdf"
    with open(p, "wb") as f:
        w.write(f)
    assert _parse_pdf(p) == ""


def test_parse_pdf_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        parse_pdf(tmp_path / "missing.pdf")


def test_parse_pdf_wrong_extension_raises(tmp_path):
    p = tmp_path / "r.txt"
    p.write_text("x")
    with pytest.raises(ValueError):
        parse_pdf(p)


def test_sample_resume_fixture_exists():
    assert (FIXTURES / "sample_resume.txt").exists()
    assert (FIXTURES / "sample_jd.txt").exists()
