"""Unit tests for the tailoring pipeline (offline path)."""
import json
from pathlib import Path

import pytest
from docx import Document

from resumetool.pipelines import tailor_resume, tailor_resume_full
from resumetool.pipelines.tailor import (
    _extract_keywords, _redact_pii, _redact_resume, _score_resume_against_jd,
    _offline_rewrite,
)
from resumetool.types import JobExperience, JobDescription, Resume


FIXTURES = Path(__file__).parent / "fixtures"


def _write_docx(tmp_path: Path, paragraphs: list[str], bullets: list[str] | None = None) -> Path:
    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    for b in bullets or []:
        d.add_paragraph(b, style="List Bullet")
    out = tmp_path / "resume.docx"
    d.save(out)
    return out


def test_tailor_resume_html_offline(tmp_path):
    resume = _write_docx(
        tmp_path,
        paragraphs=[
            "Jane Smith",
            "Senior Software Engineer with 8 years of experience on AWS and Kubernetes",
        ],
        bullets=[
            "Built distributed systems on AWS",
            "Led migration to Kubernetes",
        ],
    )
    jd = tmp_path / "jd.txt"
    jd.write_text("Looking for a Senior Cloud Engineer with AWS, Kubernetes, and Python experience.")

    out = tailor_resume(resume, jd, output_format="html", output=str(tmp_path / "out.html"))
    assert Path(out).exists()
    html = Path(out).read_text()
    assert "Jane Smith" in html
    assert "Built distributed systems" in html


def test_tailor_resume_docx_offline(tmp_path):
    resume = _write_docx(tmp_path, ["Jane Smith", "SWE with Python experience"], ["Built X"])
    jd = tmp_path / "jd.txt"
    jd.write_text("Senior Software Engineer role requiring Python, AWS, Docker.")

    out = tailor_resume(resume, jd, output_format="docx", output=str(tmp_path / "out.docx"))
    assert Path(out).exists()
    assert Path(out).stat().st_size > 1000


def test_tailor_resume_full_returns_artifacts(tmp_path):
    resume = _write_docx(
        tmp_path,
        [
            "Jane Smith",
            "Software engineer with Python and AWS experience",
            "Skills: Python, AWS, Docker",
        ],
        ["Built Python service on AWS"],
    )
    jd = tmp_path / "jd.txt"
    jd.write_text("Looking for a Python developer with AWS experience.")

    res = tailor_resume_full(
        resume, jd, output_format="html", output=str(tmp_path / "out.html"),
    )
    assert res.score > 0.0
    assert res.match.notes
    assert res.gap_analysis.suggestions
    assert res.tailored_resume.full_name
    assert "parse" in res.timings_ms
    assert "render" in res.timings_ms
    assert res.output_path.endswith(".html")


def test_tailor_resume_pdf_routes_via_html(tmp_path):
    resume = _write_docx(tmp_path, ["Jane", "SWE"], ["Did X"])
    jd = tmp_path / "jd.txt"
    jd.write_text("Looking for a Python developer with AWS experience.")
    out = tailor_resume(resume, jd, output_format="pdf", output=str(tmp_path / "out.pdf"))
    assert Path(out).exists()
    assert out.endswith(".pdf")
    # The HTML intermediate should exist alongside
    assert (tmp_path / "out.html").exists()


def test_tailor_resume_accepts_structured_json_resume(tmp_path):
    resume = tmp_path / "resume.json"
    resume.write_text(json.dumps({
        "full_name": "Jane",
        "contact_info_redacted": True,
        "summary": "Experienced Python engineer",
        "skills": ["Python", "AWS", "Kubernetes"],
        "experience": [
            {
                "company": "Acme",
                "title": "Staff Engineer",
                "duration": "2020",
                "bullet_points": ["Led migration to Kubernetes"],
            }
        ],
        "education": [],
        "certifications": [],
    }))
    jd = tmp_path / "jd.txt"
    jd.write_text("Looking for a Python Kubernetes engineer.")
    out = tailor_resume(resume, jd, output_format="html", output=str(tmp_path / "out.html"))
    assert Path(out).exists()
    assert "Jane" in Path(out).read_text()


def test_tailor_resume_rejects_unknown_format(tmp_path):
    resume = _write_docx(tmp_path, ["Jane"], ["X"])
    jd = tmp_path / "jd.txt"
    jd.write_text("X")
    with pytest.raises(ValueError, match="Unsupported output format"):
        tailor_resume(resume, jd, output_format="rtf", output=str(tmp_path / "out.rtf"))


def test_redact_pii_strips_email_and_phone():
    s = "Email me at jane.doe@gmail.com or call (555) 123-4567"
    out = _redact_pii(s)
    assert "@" not in out
    assert "555" not in out
    assert "[" in out


def test_redact_resume_keeps_name_strips_contact():
    r = Resume(
        full_name="Jane Smith",
        contact_info_redacted=True,
        summary="Reach jane@example.com",
        skills=["Python"],
        experience=[JobExperience(
            company="Acme", title="SWE", duration="2020",
            bullet_points=["Call 555-123-4569 anytime"],
        )],
        education=["jane@school.edu"],
    )
    red = _redact_resume(r)
    # Name is preserved (candidate shares their own resume with the LLM)
    assert red.full_name == "Jane Smith"
    # Email and phone are stripped
    assert "jane@example.com" not in (red.summary or "")
    assert "555-123" not in red.experience[0].bullet_points[0]
    assert "jane@" not in red.education[0]


def test_score_resume_against_jd_basic():
    r = Resume(
        full_name=None, contact_info_redacted=True,
        summary="", skills=["Python", "AWS", "Kubernetes"],
        experience=[JobExperience(company="Acme", title="SWE", duration="2020", bullet_points=["Built X"])],
    )
    jd = JobDescription(text="Need Python and AWS engineer")
    match = _score_resume_against_jd(r, jd)
    assert 0.0 < match.score <= 1.0
    assert "python" in (match.notes or "").lower() or match.score > 0


def test_score_empty_jd_returns_zero():
    r = Resume(full_name=None, contact_info_redacted=True)
    jd = JobDescription(text="")
    match = _score_resume_against_jd(r, jd)
    assert match.score == 0.0


def test_extract_keywords_filters_stopwords():
    kws = _extract_keywords("The engineer uses Python and AWS for cloud systems")
    assert "the" not in kws
    assert "engineer" in kws
    assert "python" in kws
    assert "aws" in kws


def test_offline_rewrite_reorders_skills_by_jd_overlap():
    r = Resume(
        full_name=None, contact_info_redacted=True, summary=None,
        skills=["PostgreSQL", "Kubernetes", "AWS"],
    )
    jd = JobDescription(text="We need a Kubernetes expert with AWS experience")
    gap_match = _score_resume_against_jd(r, jd)
    from resumetool.llm.guards import GapAnalysis
    gap = GapAnalysis(missing_keywords=gap_match.missing_keywords, suggestions=[])
    out = _offline_rewrite(r, jd, gap)
    assert out.skills[0] in {"Kubernetes", "Aws", "PostgreSQL"}
    # Skills that appear in JD text should rank higher
    assert {s.lower() for s in out.skills[:2]} & {"kubernetes", "aws"}


def test_tailor_resume_sample_fixtures(tmp_path):
    """Smoke test using the checked-in sample fixtures."""
    resume = FIXTURES / "sample_resume.txt"
    jd = FIXTURES / "sample_jd.txt"
    out = tailor_resume(resume, jd, output_format="html", output=str(tmp_path / "out.html"))
    assert Path(out).exists()
    html = Path(out).read_text()
    # Candidate name should appear in output
    assert "Jane Smith" in html
    # PII should be redacted
    assert "jane.smith@example.com" not in html
    assert "(555) 123-4567" not in html
